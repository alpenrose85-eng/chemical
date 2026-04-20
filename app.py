import streamlit as st
import pandas as pd
from docx import Document
import json
import os
from datetime import datetime
import io
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import re
from difflib import SequenceMatcher

class SampleNameMatcher:
    def __init__(self):
        self.surface_types = {
            'ЭПК': ['ЭПК'],
            'ШПП': ['ШПП'],
            'ПС КШ': ['ПС КШ', 'ПТ КШ', 'труба_ПТКМ', 'труба ПТКМ', 'ПТКМ', 'труба'],
            'КПП ВД': ['КПП ВД', 'ВД'],
            'КПП НД-1': ['КПП НД-1', 'КПП НД-I', 'НД-1', 'НД-I'],
            'КПП НД-2': ['КПП НД-2', 'КПП НД-II', 'НД-2', 'НД-II', 'КПП НД-IIст', 'НД-IIст']
        }
        self.letters = ['А', 'Б', 'В', 'Г']

    def parse_correct_names(self, file_content):
        """Парсинг файла с правильными названиями образцов из таблицы"""
        try:
            doc = Document(io.BytesIO(file_content))
            correct_names = []
            
            # Пытаемся сначала найти таблицы
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        number_cell = row.cells[0].text.strip()
                        name_cell = row.cells[1].text.strip()
                        if number_cell and name_cell and number_cell.isdigit():
                            correct_names.append({
                                'number': int(number_cell),
                                'original': name_cell,
                                'surface_type': self.extract_surface_type(name_cell),
                                'tube_number': self.extract_tube_number_from_correct(name_cell),
                                'letter': self.extract_letter(name_cell)
                            })
            
            # Если в таблице не нашли, проверяем параграфы
            if not correct_names:
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    # Ищем паттерн: число, затем название
                    match = re.match(r'^\s*(\d+)\s+([^\s].*)$', text)
                    if match:
                        number = match.group(1)
                        name = match.group(2).strip()
                        if number.isdigit():
                            correct_names.append({
                                'number': int(number),
                                'original': name,
                                'surface_type': self.extract_surface_type(name),
                                'tube_number': self.extract_tube_number_from_correct(name),
                                'letter': self.extract_letter(name)
                            })
            
            # Сортируем по номеру
            correct_names.sort(key=lambda x: x['number'])
            return correct_names
        except Exception as e:
            st.error(f"Ошибка при парсинге файла с правильными названиями: {str(e)}")
            return []

    def extract_tube_number_from_correct(self, correct_name):
        """Извлечение номера трубы из правильного названия"""
        # Ищем числа в скобках
        matches = re.findall(r'\((\d+)\)', correct_name)
        if matches:
            return matches[0]
        
        # Ищем числа после пробела или дефиса
        matches = re.findall(r'\s+(\d+)\)', correct_name)
        if matches:
            return matches[0]
        
        # Ищем просто числа
        matches = re.findall(r'\b(\d+)\b', correct_name)
        if matches:
            return matches[-1]  # Берем последнее число
        
        return None

    def extract_surface_type(self, name):
        """Извлечение типа поверхности нагрева из названия"""
        normalized_name = self.normalize_roman_numerals(name)
        for surface_type, patterns in self.surface_types.items():
            for pattern in patterns:
                normalized_pattern = self.normalize_roman_numerals(pattern)
                if normalized_pattern in normalized_name:
                    return surface_type
        
        # Попробуем найти частичное совпадение
        for surface_type, patterns in self.surface_types.items():
            for pattern in patterns:
                normalized_pattern = self.normalize_roman_numerals(pattern)
                if self.similar(normalized_pattern, normalized_name) > 0.7:
                    return surface_type
        
        return None

    def normalize_roman_numerals(self, text):
        """Нормализация римских цифр и суффиксов в тексте"""
        replacements = [
            (' НД-I', ' НД-1'),
            (' НД-II', ' НД-2'),
            (' НД-I ', ' НД-1 '),
            (' НД-II ', ' НД-2 '),
            ('КПП НД-I', 'КПП НД-1'),
            ('КПП НД-II', 'КПП НД-2'),
            ('НД-I', 'НД-1'),
            ('НД-II', 'НД-2'),
            ('I', '1'),
            ('II', '2'),
            ('IIст', 'II'),
            ('Iст', 'I'),
            ('-IIст', '-II'),
            ('-Iст', '-I')
        ]
        result = text
        for roman, arabic in replacements:
            result = result.replace(roman, arabic)
        return result

    def similar(self, a, b):
        """Вычисление схожести строк"""
        return SequenceMatcher(None, a, b).ratio()

    def extract_letter(self, name):
        """Извлечение буквы (А, Б, В, Г) из названия"""
        matches = re.findall(r'\([^)]*([А-Г])\)', name)
        if matches:
            return matches[0]
        
        matches = re.findall(r',\s*([А-Г])\)', name)
        if matches:
            return matches[0]
        
        matches = re.findall(r'\(([А-Г])\)', name)
        if matches:
            return matches[0]
        
        return None

    def extract_tube_number_from_protocol(self, sample_name):
        """Извлечение номера трубы из названия протокола"""
        # Ищем паттерны типа "тр. №58", "тр.61", "тр 240"
        patterns = [
            r'тр\.\s*№?\s*(\d+)',  # тр. №58, тр.61
            r'тр\s*(\d+)',         # тр 240
            r'труба\s*(\d+)',      # труба 26
            r'тр\.\s*(\d+)',       # тр. 157
            r'\((\d+)\)',          # (58)
        ]
        
        for pattern in patterns:
            match = re.search(pattern, sample_name)
            if match:
                return match.group(1)
        
        # Ищем просто числа в названии
        numbers = re.findall(r'\b\d+\b', sample_name)
        if numbers:
            # Берем наибольшее число (предполагаем, что это номер трубы)
            return max(numbers, key=lambda x: int(x))
        
        return None

    def parse_protocol_sample_name(self, sample_name):
        """Парсинг названия образца из протокола химического анализа"""
        original_name = sample_name
        
        # 1. Определяем букву нитки
        letter = None
        letter_map = {'НА': 'А', 'НБ': 'Б', 'НВ': 'В', 'НГ': 'Г', 'Н-Г': 'Г'}
        for prefix, mapped_letter in letter_map.items():
            if prefix in sample_name:
                letter = mapped_letter
                break
        
        if not letter:
            patterns = [
                r'Н[_\s\-]?([А-Г])',
                r'Н([А-Г])[_\s]',
                r'[_\s]Н([А-Г])',
            ]
            for pattern in patterns:
                matches = re.findall(pattern, sample_name)
                if matches:
                    letter = matches[0]
                    break
        
        # 2. Определяем номер трубы
        tube_number = self.extract_tube_number_from_protocol(sample_name)
        
        # 3. Определяем тип поверхности
        surface_type = self.extract_surface_type(sample_name)
        
        return {
            'original': original_name,
            'surface_type': surface_type,
            'tube_number': tube_number,
            'letter': letter
        }

    def match_samples(self, protocol_samples, correct_samples):
        """Многоэтапное сопоставление образцов по номеру трубы"""
        matched_samples = []
        unmatched_protocol = protocol_samples.copy()
        used_correct = set()
        
        # Этап 1: Совпадение по номеру трубы и типу поверхности
        matches_stage1 = self._match_by_tube_and_type(unmatched_protocol, correct_samples, used_correct)
        matched_samples.extend(matches_stage1)
        unmatched_protocol = [s for s in unmatched_protocol if s not in [m[0] for m in matches_stage1]]
        
        # Этап 2: Совпадение только по номеру трубы
        matches_stage2 = self._match_by_tube_only(unmatched_protocol, correct_samples, used_correct)
        matched_samples.extend(matches_stage2)
        unmatched_protocol = [s for s in unmatched_protocol if s not in [m[0] for m in matches_stage2]]
        
        return matched_samples, unmatched_protocol

    def _match_by_tube_and_type(self, protocol_samples, correct_samples, used_correct):
        """Сопоставление по номеру трубы и типу поверхности"""
        matches = []
        for protocol in protocol_samples:
            protocol_info = self.parse_protocol_sample_name(protocol['name'])
            for correct in correct_samples:
                if correct['original'] in used_correct:
                    continue
                
                # Сравниваем номера труб
                if (protocol_info['tube_number'] and correct['tube_number'] and 
                    protocol_info['tube_number'] == correct['tube_number']):
                    
                    # Проверяем совпадение типа поверхности
                    if (protocol_info['surface_type'] and correct['surface_type'] and
                        protocol_info['surface_type'] == correct['surface_type']):
                        matches.append((protocol, correct, "совпадение по трубе и типу"))
                        used_correct.add(correct['original'])
                        break
        return matches

    def _match_by_tube_only(self, protocol_samples, correct_samples, used_correct):
        """Сопоставление только по номеру трубы"""
        matches = []
        for protocol in protocol_samples:
            protocol_info = self.parse_protocol_sample_name(protocol['name'])
            for correct in correct_samples:
                if correct['original'] in used_correct:
                    continue
                
                # Сравниваем только номера труб
                if (protocol_info['tube_number'] and correct['tube_number'] and 
                    protocol_info['tube_number'] == correct['tube_number']):
                    matches.append((protocol, correct, "совпадение по трубе"))
                    used_correct.add(correct['original'])
                    break
        return matches


class ChemicalAnalyzer:
    def __init__(self):
        self.load_standards()
        self.name_matcher = SampleNameMatcher()
        self.all_elements = ["C", "Si", "Mn", "P", "S", "Cr", "Mo", "Ni",
                             "Cu", "Al", "Co", "Nb", "Ti", "V", "W", "Fe"]

    def load_standards(self):
        self.standards = {
            "12Х1МФ": {
                "C": (0.10, 0.15), "Si": (0.17, 0.37), "Mn": (0.40, 0.70),
                "Cr": (0.90, 1.20), "Mo": (0.25, 0.35), "V": (0.15, 0.30),
                "Ni": (None, 0.25), "Cu": (None, 0.20), "S": (None, 0.025),
                "P": (None, 0.025), "source": "ТУ 14-3Р-55-2001"
            },
            "12Х18Н12Т": {
                "C": (None, 0.12), "Si": (None, 0.80), "Mn": (1.00, 2.00),
                "Cr": (17.00, 19.00), "Ni": (11.00, 13.00), "Ti": (None, 0.70),
                "Cu": (None, 0.30), "S": (None, 0.020), "P": (None, 0.035),
                "source": "ТУ 14-3Р-55-2001"
            },
            "20": {
                "C": (0.17, 0.24), "Si": (0.17, 0.37), "Mn": (0.35, 0.65),
                "Cr": (None, 0.25), "Ni": (None, 0.25), "Cu": (None, 0.30),
                "P": (None, 0.030), "S": (None, 0.025), "source": "ТУ 14-3Р-55-2001"
            },
            "Ди82": {
                "C": (0.08, 0.12), "Si": (None, 0.5), "Mn": (0.30, 0.60),
                "Cr": (8.60, 10.00), "Ni": (None, 0.70), "Mo": (0.60, 0.80),
                "V": (0.10, 0.20), "Nb": (0.10, 0.20), "Cu": (None, 0.30),
                "S": (None, 0.015), "P": (None, 0.03), "source": "ТУ 14-3Р-55-2001"
            },
            "Ди59": {
                "C": (0.06, 0.10), "Si": (1.8, 2.2), "Mn": (12.00, 13.50),
                "Cr": (11.50, 13.00), "Ni": (1.8, 2.5), "Nb": (0.60, 1.00),
                "Cu": (2.00, 2.50), "S": (None, 0.02), "P": (None, 0.03),
                "source": "ТУ 14-3Р-55-2001"
            }
        }
        if os.path.exists("user_standards.json"):
            with open("user_standards.json", "r", encoding="utf-8") as f:
                user_std = json.load(f)
                self.standards.update(user_std)

    def save_user_standards(self):
        with open("user_standards.json", "w", encoding="utf-8") as f:
            predefined = ["12Х1МФ", "12Х18Н12Т", "20", "Ди82", "Ди59"]
            user_standards = {k: v for k, v in self.standards.items() if k not in predefined}
            json.dump(user_standards, f, ensure_ascii=False, indent=2)

    def parse_protocol_file(self, file_content):
        try:
            doc = Document(io.BytesIO(file_content))
            samples = []
            current_sample = None
            
            # Парсим параграфы для получения информации об образцах
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                if "Наименование образца:" in text:
                    sample_name = text.split("Наименование образца:")[1].strip()
                    current_sample = {
                        "name": sample_name,
                        "steel_grade": None,
                        "composition": {},
                        "original_name": sample_name
                    }
                    samples.append(current_sample)
                
                elif "Химический состав металла образца соответствует марке стали:" in text:
                    if current_sample:
                        grade_text = text.split("марке стали:")[1].strip()
                        # Убираем звездочки и дополнительные комментарии
                        grade_text = re.sub(r'\*+', '', grade_text).strip()
                        grade_text = grade_text.split(',')[0].strip()
                        current_sample["steel_grade"] = grade_text
            
            # Парсим таблицы для получения состава
            table_index = 0
            for table in doc.tables:
                if table_index < len(samples):
                    composition = self.parse_composition_table(table)
                    samples[table_index]["composition"] = composition
                    table_index += 1
            
            return samples
        except Exception as e:
            st.error(f"Ошибка при парсинге файла: {str(e)}")
            return []

    def parse_composition_table(self, table):
        """Парсинг таблицы состава - ВОЗВРАЩАЕМ ИСХОДНЫЙ РАБОЧИЙ ПАРСИНГ"""
        composition = {}
        try:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if len(table_data) < 13:
                st.warning(f"Таблица имеет только {len(table_data)} строк, ожидалось минимум 13")
                return composition

            # Берем данные из строк как в исходной рабочей версии
            headers_row1 = table_data[0]
            values_row1 = table_data[5]
            headers_row2 = table_data[7]
            values_row2 = table_data[12]

            for i, header in enumerate(headers_row1):
                if header in self.all_elements and i < len(values_row1):
                    value_str = values_row1[i]
                    try:
                        value_str = value_str.replace(',', '.').replace(' ', '')
                        if '±' in value_str:
                            value_str = value_str.split('±')[0]
                        value = float(value_str)
                        composition[header] = value
                    except (ValueError, IndexError):
                        continue

            for i, header in enumerate(headers_row2):
                if header in self.all_elements and i < len(values_row2):
                    value_str = values_row2[i]
                    try:
                        value_str = value_str.replace(',', '.').replace(' ', '')
                        if '±' in value_str:
                            value_str = value_str.split('±')[0]
                        value = float(value_str)
                        composition[header] = value
                    except (ValueError, IndexError):
                        continue

            return composition
        except Exception as e:
            st.error(f"Ошибка при парсинге таблицы: {str(e)}")
            return {}

    def match_sample_names(self, samples, correct_names_file):
        """Сопоставление названий образцов"""
        if not correct_names_file:
            return samples, []
        
        correct_samples = self.name_matcher.parse_correct_names(correct_names_file.getvalue())
        if not correct_samples:
            st.warning("Не удалось загрузить правильные названия образцов")
            return samples, []
        
        matched_pairs, unmatched_protocol = self.name_matcher.match_samples(samples, correct_samples)
        
        matched_samples = []
        for protocol_sample, correct_sample, match_stage in matched_pairs:
            corrected_sample = protocol_sample.copy()
            corrected_sample['original_name'] = protocol_sample['name']
            corrected_sample['name'] = correct_sample['original']
            corrected_sample['correct_number'] = correct_sample['number']
            corrected_sample['automatically_matched'] = True
            corrected_sample['match_stage'] = match_stage
            matched_samples.append(corrected_sample)
        
        unmatched_samples = []
        for sample in unmatched_protocol:
            sample['original_name'] = sample['name']
            sample['correct_number'] = None
            sample['automatically_matched'] = False
            unmatched_samples.append(sample)
        
        # Объединяем и сортируем по correct_number
        all_samples = matched_samples + unmatched_samples
        
        if matched_samples:
            st.success(f"✅ Автоматически сопоставлено {len(matched_samples)} образцов")
            with st.expander("📋 Детали автоматического сопоставления"):
                match_data = []
                for sample in matched_samples:
                    protocol_info = self.name_matcher.parse_protocol_sample_name(sample['original_name'])
                    match_data.append({
                        'Номер': sample['correct_number'],
                        'Исходное название': sample['original_name'],
                        'Правильное название': sample['name'],
                        'Этап': sample.get('match_stage', 'н/д'),
                        'Тип': protocol_info['surface_type'] or 'н/д',
                        'Труба': protocol_info['tube_number'] or 'н/д'
                    })
                match_data.sort(key=lambda x: x['Номер'])
                st.table(pd.DataFrame(match_data))
        
        if unmatched_samples:
            st.warning(f"⚠️ Не удалось автоматически сопоставить {len(unmatched_samples)} образцов")
            with st.expander("🔍 Просмотр несопоставленных образцов"):
                unmatched_data = []
                for sample in unmatched_samples:
                    protocol_info = self.name_matcher.parse_protocol_sample_name(sample['name'])
                    unmatched_data.append({
                        'Образец': sample['original_name'],
                        'Марка стали': sample['steel_grade'],
                        'Тип': protocol_info['surface_type'] or 'н/д',
                        'Труба': protocol_info['tube_number'] or 'н/д'
                    })
                st.table(pd.DataFrame(unmatched_data))
        
        return all_samples, correct_samples

    def apply_manual_matches(self, samples, correct_dict, manual_matches):
        """Применение ручных сопоставлений к образцам"""
        updated_samples = []
        
        # Определяем конфликты (когда одно правильное название выбрано для нескольких образцов)
        used_names = {}
        conflicts = []
        
        for original_name, correct_name in manual_matches.items():
            if correct_name:
                if correct_name in used_names:
                    conflicts.append((original_name, correct_name, used_names[correct_name]))
                else:
                    used_names[correct_name] = original_name
        
        # Разрешаем конфликты - оставляем только последнее сопоставление
        resolved_matches = {}
        for original_name, correct_name in manual_matches.items():
            if correct_name:
                if correct_name not in resolved_matches.values():
                    resolved_matches[original_name] = correct_name
                else:
                    # Находим, какой образец уже имеет это название
                    existing_original = [k for k, v in resolved_matches.items() if v == correct_name][0]
                    # Удаляем у старого образца
                    st.warning(f"Название '{correct_name}' выбрано для нескольких образцов. Оставлено для '{original_name}', образец '{existing_original}' останется без названия.")
        
        # Обновляем образцы с учетом ручных сопоставлений
        for sample in samples:
            updated_sample = sample.copy()
            
            if sample['original_name'] in resolved_matches:
                correct_name = resolved_matches[sample['original_name']]
                
                if correct_name and correct_name in correct_dict:
                    # Образец сопоставлен
                    updated_sample['name'] = correct_name
                    updated_sample['correct_number'] = correct_dict[correct_name]['number']
                    updated_sample['manually_matched'] = True
                    updated_sample['automatically_matched'] = False
                    updated_sample['match_stage'] = "ручное сопоставление"
                else:
                    # Образец не сопоставлен (некорректное название)
                    updated_sample['name'] = sample['original_name']
                    updated_sample['correct_number'] = None
                    updated_sample['manually_matched'] = False
                    updated_sample['automatically_matched'] = False
            else:
                # Если нет ручного сопоставления, оставляем как было
                if not sample.get('automatically_matched'):
                    updated_sample['name'] = sample['original_name']
                    updated_sample['correct_number'] = None
                    updated_sample['manually_matched'] = False
                    updated_sample['automatically_matched'] = False
                else:
                    updated_sample['manually_matched'] = False
            
            updated_samples.append(updated_sample)
        
        return updated_samples

    def add_manual_matching_interface(self, samples, correct_samples):
        """Интерфейс для ручного сопоставления образцов"""
        st.header("🔧 Ручное сопоставление образцов")
        
        # Инициализируем session_state для хранения ручных сопоставлений
        if 'manual_matches' not in st.session_state:
            st.session_state.manual_matches = {}
        
        # Создаем словарь для быстрого поиска правильных названий
        correct_dict = {cs['original']: cs for cs in correct_samples}
        correct_names_list = [cs['original'] for cs in correct_samples]
        
        # Группируем образцы по марке стали
        samples_by_grade = {}
        for sample in samples:
            grade = sample.get('steel_grade', 'Не указана')
            if grade not in samples_by_grade:
                samples_by_grade[grade] = []
            samples_by_grade[grade].append(sample)
        
        # Создаем интерфейс для ручного сопоставления
        for grade, grade_samples in samples_by_grade.items():
            st.subheader(f"Марка стали: {grade}")
            
            for i, sample in enumerate(grade_samples):
                col1, col2 = st.columns([2, 3])
                
                with col1:
                    # Отображаем информацию об образце
                    st.write(f"**{sample['original_name']}**")
                    if sample.get('steel_grade'):
                        st.write(f"*Марка: {sample['steel_grade']}*")
                    
                    protocol_info = self.name_matcher.parse_protocol_sample_name(sample['original_name'])
                    if protocol_info['tube_number']:
                        st.write(f"*Труба: {protocol_info['tube_number']}*")
                    
                    # Показываем текущий статус
                    current_status = ""
                    if sample.get('automatically_matched'):
                        current_status = "✅ Автоматически сопоставлен"
                    elif sample['original_name'] in st.session_state.manual_matches:
                        current_status = "📝 Ручное сопоставление"
                    else:
                        current_status = "❌ Не сопоставлен"
                    
                    st.write(f"*Статус: {current_status}*")
                
                with col2:
                    # Создаем список всех доступных правильных названий
                    all_options = ["Не сопоставлен"] + correct_names_list
                    
                    # Определяем текущее значение для этого образца
                    current_value = st.session_state.manual_matches.get(
                        sample['original_name'], 
                        sample['name'] if sample.get('automatically_matched') else "Не сопоставлен"
                    )
                    
                    # Если текущее значение не в списке, сбрасываем
                    if current_value not in all_options:
                        current_value = "Не сопоставлен"
                    
                    # Создаем selectbox со всеми вариантами
                    selected = st.selectbox(
                        f"Выберите правильное название для образца {i+1}",
                        options=all_options,
                        index=all_options.index(current_value) if current_value in all_options else 0,
                        key=f"manual_match_{sample['original_name']}_{grade}_{i}"
                    )
                    
                    # Сохраняем выбор в session_state
                    if selected != "Не сопоставлен":
                        st.session_state.manual_matches[sample['original_name']] = selected
                    elif sample['original_name'] in st.session_state.manual_matches:
                        del st.session_state.manual_matches[sample['original_name']]
            
            st.markdown("---")
        
        # Кнопки управления
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔄 Сбросить все ручные сопоставления"):
                st.session_state.manual_matches = {}
                st.rerun()
        
        with col2:
            if st.button("✅ Применить ручное сопоставление"):
                # Применяем ручные сопоставления
                updated_samples = self.apply_manual_matches(samples, correct_dict, st.session_state.manual_matches)
                
                # Показываем сводку изменений
                st.success(f"✅ Ручное сопоставление применено! Обновлено {len(st.session_state.manual_matches)} образцов.")
                
                # Показываем изменения
                with st.expander("📋 Сводка изменений"):
                    changes = []
                    for sample in updated_samples:
                        original_sample = next((s for s in samples if s['original_name'] == sample['original_name']), None)
                        if original_sample:
                            if sample.get('manually_matched') and original_sample.get('automatically_matched'):
                                changes.append({
                                    'Образец': sample['original_name'],
                                    'Было': original_sample['name'],
                                    'Стало': sample['name'],
                                    'Тип': 'Переназначение'
                                })
                            elif sample.get('manually_matched') and not original_sample.get('automatically_matched'):
                                changes.append({
                                    'Образец': sample['original_name'],
                                    'Было': 'Не сопоставлен',
                                    'Стало': sample['name'],
                                    'Тип': 'Новое сопоставление'
                                })
                            elif original_sample.get('automatically_matched') and not sample.get('automatically_matched'):
                                changes.append({
                                    'Образец': sample['original_name'],
                                    'Было': original_sample['name'],
                                    'Стало': 'Не сопоставлен',
                                    'Тип': 'Удалено сопоставление'
                                })
                    
                    if changes:
                        st.table(pd.DataFrame(changes))
                    else:
                        st.info("Изменений нет")
                
                # Обновляем session_state и перезапускаем
                st.session_state.samples = updated_samples
                st.rerun()
        
        return samples

    def check_element_compliance(self, element, value, standard):
        """Проверка соответствия элемента нормам"""
        if element not in standard or element == "source":
            return "normal"
        
        min_val, max_val = standard[element]
        
        if min_val is not None and value < min_val:
            return "deviation"
        elif max_val is not None and value > max_val:
            return "deviation"
        else:
            return "normal"

    def create_report_tables(self, samples):
        """Создание таблиц отчета с правильной нумерацией - ТОЛЬКО СОПОСТАВЛЕННЫЕ ОБРАЗЦЫ"""
        if not samples:
            return None
        
        # Применяем текущие ручные сопоставления из session_state
        if 'manual_matches' in st.session_state and st.session_state.manual_matches:
            # Получаем правильные названия из сохраненных образцов
            correct_samples = st.session_state.get('correct_samples', [])
            if correct_samples:
                correct_dict = {cs['original']: cs for cs in correct_samples}
                samples = self.apply_manual_matches(samples, correct_dict, st.session_state.manual_matches)
        
        # Фильтруем только сопоставленные образцы (те, у которых есть correct_number)
        matched_samples = [s for s in samples if s.get('correct_number') is not None]
        
        if not matched_samples:
            st.warning("❌ Нет сопоставленных образцов для создания таблиц")
            return None
        
        # Группируем сопоставленные образцы по марке стали
        steel_grades = list(set(sample["steel_grade"] for sample in matched_samples if sample["steel_grade"]))
        tables = {}
        
        for grade in steel_grades:
            grade_samples = [s for s in matched_samples if s["steel_grade"] == grade]
            
            if grade not in self.standards:
                st.warning(f"Нет нормативов для марки стали: {grade}")
                continue
            
            standard = self.standards[grade]
            
            # Определяем порядок элементов для этой марки стали
            if grade == "12Х1МФ":
                main_elements = ["C", "Si", "Mn", "Cr", "Mo", "V", "Ni"]
                harmful_elements = ["Cu", "S", "P"]
                other_elements = [elem for elem in standard.keys() 
                                 if elem not in main_elements + harmful_elements + ["source"]]
                norm_elements = main_elements + other_elements + harmful_elements
            elif grade == "20":
                main_elements = ["C", "Si", "Mn"]
                harmful_elements = ["P", "S"]
                other_elements = [elem for elem in standard.keys() 
                                 if elem not in main_elements + harmful_elements + ["source"]]
                norm_elements = main_elements + other_elements + harmful_elements
            else:
                norm_elements = [elem for elem in standard.keys() if elem != "source"]
            
            # Сортируем образцы по correct_number (порядку из файла с правильными названиями)
            sorted_samples = sorted(
                grade_samples,
                key=lambda x: x.get('correct_number', float('inf'))
            )
            
            # Создаем данные для таблицы
            data = []
            compliance_data = []
            
            # Добавляем образцы с нумерацией с 1
            for idx, sample in enumerate(sorted_samples, 1):
                row = {
                    "№": idx,
                    "Образец": sample["name"]
                }
                compliance_row = {"№": "normal", "Образец": "normal"}
                
                # Добавляем элементы состава
                for elem in norm_elements:
                    if elem in sample["composition"]:
                        value = sample["composition"][elem]
                        
                        # Форматируем значение в зависимости от элемента
                        if elem in ["S", "P"]:
                            row[elem] = f"{value:.3f}".replace('.', ',')
                        else:
                            row[elem] = f"{value:.2f}".replace('.', ',')
                        
                        # Проверяем соответствие нормам
                        status = self.check_element_compliance(elem, value, standard)
                        compliance_row[elem] = status
                    else:
                        row[elem] = "-"
                        compliance_row[elem] = "normal"
                
                data.append(row)
                compliance_data.append(compliance_row)
            
            # Добавляем строку требований
            requirements_row = {"№": "", "Образец": f"Требования ТУ 14-3Р-55-2001 для стали марки {grade}"}
            requirements_compliance = {"№": "requirements", "Образец": "requirements"}
            
            for elem in norm_elements:
                if elem in standard:
                    min_val, max_val = standard[elem]
                    if min_val is not None and max_val is not None:
                        if elem in ["S", "P"]:
                            requirements_row[elem] = f"{min_val:.3f}-{max_val:.3f}".replace('.', ',')
                        else:
                            requirements_row[elem] = f"{min_val:.2f}-{max_val:.2f}".replace('.', ',')
                    elif min_val is not None:
                        if elem in ["S", "P"]:
                            requirements_row[elem] = f"≥{min_val:.3f}".replace('.', ',')
                        else:
                            requirements_row[elem] = f"≥{min_val:.2f}".replace('.', ',')
                    elif max_val is not None:
                        if elem in ["S", "P"]:
                            requirements_row[elem] = f"≤{max_val:.3f}".replace('.', ',')
                        else:
                            requirements_row[elem] = f"≤{max_val:.2f}".replace('.', ',')
                    else:
                        requirements_row[elem] = "не нормируется"
                else:
                    requirements_row[elem] = "-"
                
                requirements_compliance[elem] = "requirements"
            
            data.append(requirements_row)
            compliance_data.append(requirements_compliance)
            
            # Создаем DataFrame
            df = pd.DataFrame(data)
            
            tables[grade] = {
                "data": df,
                "compliance": compliance_data,
                "samples": sorted_samples,
                "requirements": requirements_row
            }
        
        return tables

    def apply_styling(self, df, compliance_data):
        """Применение стилей к таблице"""
        styled = df.style
        
        # Применяем стили для каждой ячейки
        for i in range(len(df)):
            for j, col in enumerate(df.columns):
                if i < len(compliance_data) and col in compliance_data[i]:
                    status = compliance_data[i][col]
                    
                    if status == "deviation":
                        styled = styled.set_properties(
                            subset=(i, col),
                            **{'background-color': '#ffcccc', 'color': '#cc0000', 'font-weight': 'bold'}
                        )
                    elif status == "requirements":
                        styled = styled.set_properties(
                            subset=(i, col),
                            **{'background-color': '#f0f0f0', 'font-style': 'italic'}
                        )
        
        return styled


def set_font_times_new_roman(doc):
    """Установка шрифта Times New Roman для всего документа"""
    styles = doc.styles
    for style in styles:
        if hasattr(style, 'font'):
            style.font.name = 'Times New Roman'
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'


def create_word_report(samples, analyzer, report_tables=None):
    """Создание Word отчета"""
    try:
        # Если таблицы не предоставлены, создаем их
        if report_tables is None:
            # Применяем текущие ручные сопоставления перед созданием таблиц
            if 'manual_matches' in st.session_state and st.session_state.manual_matches:
                correct_samples = st.session_state.get('correct_samples', [])
                if correct_samples:
                    correct_dict = {cs['original']: cs for cs in correct_samples}
                    samples = analyzer.apply_manual_matches(samples, correct_dict, st.session_state.manual_matches)
            
            # Создаем таблицы
            report_tables = analyzer.create_report_tables(samples)
            if not report_tables:
                st.warning("Нет данных для создания отчета")
                return
        
        doc = Document()
        set_font_times_new_roman(doc)
        
        # Заголовок
        title = doc.add_heading('Протокол анализа химического состава', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        
        # Считаем только сопоставленные образцы
        matched_samples = [s for s in samples if s.get('correct_number') is not None]
        doc.add_paragraph(f"Проанализировано образцов: {len(matched_samples)}")
        doc.add_paragraph("")
        
        # Легенда
        doc.add_heading('Легенда', level=1)
        legend_table = doc.add_table(rows=3, cols=2)
        legend_table.style = 'Table Grid'
        legend_table.cell(0, 0).text = "Цвет"
        legend_table.cell(0, 1).text = "Значение"
        legend_table.cell(1, 0).text = "🔴"
        legend_table.cell(1, 1).text = "Отклонение от норм"
        legend_table.cell(2, 0).text = "⚪"
        legend_table.cell(2, 1).text = "Нормативные требования"
        
        doc.add_paragraph()
        
        # Добавляем таблицы для каждой марки стали
        for grade, table_data in report_tables.items():
            doc.add_heading(f'Марка стали: {grade}', level=1)
            
            df = table_data["data"]
            word_table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
            word_table.style = 'Table Grid'
            
            # Заголовки
            for j, col in enumerate(df.columns):
                word_table.cell(0, j).text = str(col)
            
            # Данные
            for i, row in df.iterrows():
                for j, col in enumerate(df.columns):
                    word_table.cell(i+1, j).text = str(row[col])
            
            doc.add_paragraph()
        
        # Сохраняем документ
        doc.save("химический_анализ_отчет.docx")
        
        # Предоставляем ссылку для скачивания
        with open("химический_анализ_отчет.docx", "rb") as file:
            btn = st.download_button(
                label="📥 Скачать отчет в формате Word",
                data=file,
                file_name=f"химический_анализ_отчет_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        st.success("✅ Отчет успешно создан! Нажмите кнопку выше для скачивания.")
        
    except Exception as e:
        st.error(f"Ошибка при создании Word отчета: {str(e)}")


def main():
    st.set_page_config(page_title="Анализатор химсостава металла", layout="wide")
    st.title("🔬 Анализатор химического состава металла")
    
    # Инициализируем анализатор
    analyzer = ChemicalAnalyzer()
    
    # Инициализируем session_state для хранения данных
    if 'samples' not in st.session_state:
        st.session_state.samples = []
    if 'correct_samples' not in st.session_state:
        st.session_state.correct_samples = []
    if 'manual_matches' not in st.session_state:
        st.session_state.manual_matches = {}
    if 'report_tables' not in st.session_state:
        st.session_state.report_tables = None
    if 'last_correct_hash' not in st.session_state:
        st.session_state.last_correct_hash = None
    if 'last_protocols_hash' not in st.session_state:
        st.session_state.last_protocols_hash = None
    if 'parsed_samples' not in st.session_state:
        st.session_state.parsed_samples = []
    
    # Сайдбар с нормативами
    with st.sidebar:
        st.header("📋 Управление нормативами")
        
        st.subheader("Существующие марки стали")
        selected_standard = st.selectbox(
            "Выберите марку для просмотра",
            options=list(analyzer.standards.keys())
        )
        
        if selected_standard:
            st.write(f"**Норматив для {selected_standard}:**")
            standard = analyzer.standards[selected_standard]
            for elem, value_range in standard.items():
                if elem == "source":
                    continue
                if isinstance(value_range, tuple) and len(value_range) == 2:
                    min_val, max_val = value_range
                    if min_val is not None and max_val is not None:
                        st.write(f"- {elem}: {min_val:.3f} - {max_val:.3f}")
                    elif min_val is not None:
                        st.write(f"- {elem}: ≥ {min_val:.3f}")
                    elif max_val is not None:
                        st.write(f"- {elem}: ≤ {max_val:.3f}")
            st.write(f"Источник: {standard.get('source', 'не указан')}")
    
    # Основная часть
    st.header("📁 Загрузка файлов")
    
    # Загрузка файла с правильными названиями
    st.subheader("1. Загрузите файл с правильными названиями образцов")
    correct_names_file = st.file_uploader(
        "Файл с правильными названиями (.docx)",
        type=["docx"],
        key="correct_names"
    )
    
    if correct_names_file:
        # Вычисляем хэш файла (имя + размер)
        current_hash = (correct_names_file.name, correct_names_file.size)
        if st.session_state.last_correct_hash != current_hash:
            st.session_state.correct_samples = analyzer.name_matcher.parse_correct_names(correct_names_file.getvalue())
            st.session_state.last_correct_hash = current_hash
            # При изменении файла сбрасываем ручные сопоставления и образцы
            st.session_state.manual_matches = {}
            st.session_state.samples = []
        
        if st.session_state.correct_samples:
            st.success(f"✅ Загружено {len(st.session_state.correct_samples)} правильных названий образцов")
            
            with st.expander("📋 Просмотр загруженных названий"):
                preview_data = []
                for sample in st.session_state.correct_samples:
                    preview_data.append({
                        'Номер': sample['number'],
                        'Название': sample['original'],
                        'Тип': sample['surface_type'] or 'н/д',
                        'Труба': sample['tube_number'] or 'н/д',
                        'Нитка': sample['letter'] or 'н/д'
                    })
                st.table(pd.DataFrame(preview_data))
    
    # Загрузка протоколов химического анализа
    st.subheader("2. Загрузите файлы протоколов химического анализа")
    uploaded_files = st.file_uploader(
        "Файлы протоколов (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="protocol_files"
    )
    
    if uploaded_files:
        # Вычисляем хэш набора файлов (имена и размеры)
        current_protocols_hash = tuple((f.name, f.size) for f in uploaded_files)
        if st.session_state.last_protocols_hash != current_protocols_hash:
            # Парсим заново
            all_samples = []
            for f in uploaded_files:
                all_samples.extend(analyzer.parse_protocol_file(f.getvalue()))
            st.session_state.parsed_samples = all_samples
            st.session_state.last_protocols_hash = current_protocols_hash
            # Сбрасываем старые сопоставления
            st.session_state.samples = []
            st.session_state.manual_matches = {}
        
        all_samples = st.session_state.parsed_samples
        
        if all_samples:
            st.success(f"✅ Загружено {len(all_samples)} образцов из протоколов")
            
            # Если есть правильные названия, проводим сопоставление (только если samples еще не заполнены)
            if correct_names_file and st.session_state.correct_samples:
                st.header("🔍 Сопоставление названий образцов")
                
                # Если samples пусты, выполняем автоматическое сопоставление
                if not st.session_state.samples:
                    # Автоматическое сопоставление
                    matched_samples, _ = analyzer.match_sample_names(
                        all_samples, 
                        correct_names_file
                    )
                    st.session_state.samples = matched_samples
                
                # Ручное сопоставление (работает с текущими st.session_state.samples)
                st.session_state.samples = analyzer.add_manual_matching_interface(
                    st.session_state.samples, 
                    st.session_state.correct_samples
                )
            
            # Сохраняем образцы в session_state (если еще не сохранены)
            if not st.session_state.samples:
                st.session_state.samples = all_samples
            
            # Отображаем результаты
            if st.session_state.samples:
                st.header("📊 Результаты анализа")
                
                # Создаем таблицы отчета - ТОЛЬКО СОПОСТАВЛЕННЫЕ ОБРАЗЦЫ
                report_tables = analyzer.create_report_tables(st.session_state.samples)
                
                if report_tables:
                    # Сохраняем таблицы в session_state
                    st.session_state.report_tables = report_tables
                    
                    # Отображаем легенду
                    st.markdown("""
                    **Легенда:**
                    - <span style='background-color: #ffcccc; padding: 2px 5px; border-radius: 3px;'>🔴 Красный</span> - отклонение от норм
                    - <span style='background-color: #f0f0f0; padding: 2px 5px; border-radius: 3px;'>⚪ Серый</span> - нормативные требования
                    """, unsafe_allow_html=True)
                    
                    # Отображаем таблицы для каждой марки стали
                    export_tables = {}
                    for grade, table_data in report_tables.items():
                        st.subheader(f"Марка стали: {grade}")
                        
                        # Применяем стили и отображаем таблицу
                        styled_table = analyzer.apply_styling(
                            table_data["data"], 
                            table_data["compliance"]
                        )
                        st.dataframe(styled_table, use_container_width=True, hide_index=True)
                        
                        # Сохраняем для экспорта
                        export_tables[grade] = table_data
                    
                    # Кнопка для создания Word отчета
                    if st.button("📄 Создать Word отчет"):
                        # Используем сохраненные таблицы из session_state
                        create_word_report(st.session_state.samples, analyzer, st.session_state.report_tables)
                else:
                    st.warning("❌ Нет сопоставленных образцов для создания таблиц отчета. Пожалуйста, выполните ручное сопоставление.")
                
                # Детальная информация об образцах
                st.header("📋 Детальная информация об образцах")
                
                # Разделяем на сопоставленные и несопоставленные
                matched_samples = [s for s in st.session_state.samples if s.get('correct_number') is not None]
                unmatched_samples = [s for s in st.session_state.samples if s.get('correct_number') is None]
                
                if matched_samples:
                    with st.expander(f"✅ Сопоставленные образцы ({len(matched_samples)} шт.)"):
                        for sample in matched_samples:
                            st.write(f"**{sample['name']}**")
                            st.write(f"  - Исходное название: {sample['original_name']}")
                            st.write(f"  - Марка стали: {sample['steel_grade']}")
                            st.write(f"  - Номер в списке: {sample['correct_number']}")
                            
                            if sample.get('composition'):
                                st.write("  - Химический состав:")
                                for element, value in sample['composition'].items():
                                    st.write(f"    - {element}: {value:.3f}")
                            
                            st.write("---")
                
                if unmatched_samples:
                    with st.expander(f"⚠️ Несопоставленные образцы ({len(unmatched_samples)} шт.)"):
                        st.info("Эти образцы не войдут в финальные таблицы отчета")
                        for sample in unmatched_samples:
                            st.write(f"**{sample['original_name']}**")
                            st.write(f"  - Марка стали: {sample['steel_grade']}")
                            
                            if sample.get('composition'):
                                st.write("  - Химический состав:")
                                for element, value in sample['composition'].items():
                                    st.write(f"    - {element}: {value:.3f}")
                            
                            st.write("---")


if __name__ == "__main__":
    main()
